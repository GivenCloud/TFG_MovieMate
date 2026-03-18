"""
Script to expand §4.3.2 Arquitectura del front-end and add §4.5.4 Kubernetes/Minikube
to Memoria TFG_rev260318.docx → saved as Memoria TFG_rev260319.docx
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import copy

# ─────────────────────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────────────────────
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _qn(tag):
    return f'{{{NS}}}{tag}'


def _new_list_num(doc, abstract_num_id_val, new_num_id_val):
    """Add a new <w:num> to the numbering part referencing abstract_num_id_val."""
    numbering_root = doc.part.numbering_part._element
    num = OxmlElement('w:num')
    num.set(_qn('numId'), str(new_num_id_val))
    absRef = OxmlElement('w:abstractNumId')
    absRef.set(_qn('val'), str(abstract_num_id_val))
    num.append(absRef)
    numbering_root.append(num)


def make_para(style_id, text, num_id=None, bold_prefix=None):
    """
    Create a <w:p> element.
    - style_id: Word style ID (e.g. 'Normal', 'Ttulo3', 'Prrafodelista')
    - text: paragraph text
    - num_id: if set, add numPr for list paragraphs
    - bold_prefix: if set, the first part of the text (up to bold_prefix length) is bold
    """
    p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(_qn('val'), style_id)
    pPr.append(pStyle)

    if num_id is not None:
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(_qn('val'), '0')
        numId_el = OxmlElement('w:numId')
        numId_el.set(_qn('val'), str(num_id))
        numPr.append(ilvl)
        numPr.append(numId_el)
        pPr.append(numPr)

    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)

    return p


def insert_sequence(ref_elem, items):
    """
    Insert a sequence of paragraphs after ref_elem.
    items = list of <w:p> elements (already created via make_para)
    Returns the last inserted element.
    """
    current = ref_elem
    for p_elem in items:
        current.addnext(p_elem)
        current = p_elem
    return current


def remove_paras(doc, start_idx, end_idx_inclusive):
    """Remove document paragraphs [start_idx .. end_idx_inclusive]."""
    body = doc.element.body
    to_remove = [doc.paragraphs[i]._element for i in range(start_idx, end_idx_inclusive + 1)]
    for elem in to_remove:
        body.remove(elem)


# ─────────────────────────────────────────────────────────────
#  Content definitions
# ─────────────────────────────────────────────────────────────

# New numIds for bullet lists (abstractNumId=15 → same bullet style as numId=39)
ABS_NUM_ID = 15
NEW_NUM_IDS = list(range(50, 63))  # 50..62

def build_frontend_section(num_ids):
    """Build the paragraphs for §4.3.2 (all subsections)."""
    nid = iter(num_ids)
    paras = []

    def p(text, style='Normal', num_id=None):
        paras.append(make_para(style, text, num_id=num_id))

    def h4(text):
        paras.append(make_para('Ttulo4', text))

    def li(text, nid_val):
        paras.append(make_para('Prrafodelista', text, num_id=nid_val))

    def figure(text):
        paras.append(make_para('Normal', text))

    # ── Intro ───────────────────────────────────────────────
    p("El front-end de MovieMate se ha implementado como una Single Page "
      "Application (SPA) desarrollada con React 19 y TypeScript. La arquitectura "
      "sigue un enfoque modular centrado en características (feature-based), "
      "organizando el código en módulos independientes que agrupan páginas, "
      "componentes, hooks y lógica de acceso a datos. A continuación se describen "
      "los principales aspectos estructurales y las páginas que conforman la "
      "aplicación.")

    # ── 4.3.2.1 Estructura del proyecto ─────────────────────
    h4("Estructura del proyecto")

    p("La raíz del código fuente se organiza en los siguientes directorios "
      "principales dentro de src/:")

    nid_val = next(nid)
    li("pages/ — vistas completas de la aplicación, cada una asociada a una ruta navegable.", nid_val)
    li("features/ — módulos por funcionalidad (listas, valoraciones, perfil, notificaciones, "
       "etc.), cada uno con sus propios componentes, hooks y lógica.", nid_val)
    li("components/ — componentes compartidos y reutilizables (tarjetas de póster, modales, "
       "botones de acción) empleados en múltiples páginas.", nid_val)
    li("hooks/ — hooks personalizados que encapsulan lógica de acceso a datos y gestión de estado.", nid_val)
    li("api/ — módulos de acceso a la API del servidor organizados por recurso: usersApi, "
       "listsApi, ratingsApi, contentApi, notificationsApi.", nid_val)
    li("stores/ — estado global sincrónico gestionado con Zustand (tienda de autenticación "
       "useAuthStore).", nid_val)
    li("lib/ — utilidades y funciones auxiliares: cliente Axios preconfigurado, jerarquía de "
       "claves de caché (queryKeys), helpers de slugificación y formato.", nid_val)
    li("types/ — definiciones de tipos TypeScript compartidas entre módulos.", nid_val)

    p("Esta organización favorece la cohesión interna de cada módulo, reduce el acoplamiento "
      "entre características y facilita la localización del código durante el mantenimiento y "
      "la evolución de la aplicación.")

    figure("[Figura 4.1: Árbol de directorios del proyecto front-end de MovieMate.]")

    # ── 4.3.2.2 Navegación y autenticación ──────────────────
    h4("Navegación y autenticación")

    p("La navegación entre páginas se gestiona mediante React Router v6. Las rutas se declaran "
      "de forma centralizada en el componente raíz App.tsx, distinguiendo entre rutas públicas "
      "(accesibles sin sesión) y rutas protegidas (que requieren autenticación). El acceso a "
      "estas últimas se controla mediante el componente guardián PrivateRoute, que comprueba el "
      "estado de autenticación almacenado en Zustand y redirige al usuario a la pantalla de "
      "inicio de sesión si no dispone de sesión activa.")

    p("El flujo de autenticación sigue los pasos siguientes: el usuario introduce sus "
      "credenciales en el formulario; la petición se envía al endpoint POST /api/auth/login; "
      "si las credenciales son válidas, el servidor devuelve un token JWT que se almacena tanto "
      "en el store de Zustand (estado en memoria) como en el almacenamiento local del navegador "
      "(persistencia entre recargas de página); a partir de ese momento, una instancia global "
      "de Axios incorpora automáticamente el token como cabecera Authorization: Bearer <token> "
      "en cada petición autenticada.")

    figure("[Figura 4.2: Diagrama del flujo de autenticación y protección de rutas en el "
           "front-end.]")

    # ── 4.3.2.3 Gestión del estado y comunicación con el servidor ───
    h4("Gestión del estado y comunicación con el servidor")

    p("La gestión del estado se divide en dos capas claramente separadas para evitar mezclar "
      "lógica de servidor y estado local en los componentes.")

    p("El estado global sincrónico, como los datos del usuario autenticado y el token JWT, se "
      "gestiona mediante Zustand. Su store reactivo de bajo peso permite que cualquier "
      "componente suscriba cambios de forma eficiente sin renderizados innecesarios.")

    p("El estado asíncrono derivado del servidor se gestiona con TanStack Query v5, que "
      "automatiza el ciclo de vida de las peticiones HTTP: almacena los resultados en caché "
      "mediante un sistema de claves estructuradas (queryKeys), ejecuta revalidaciones "
      "transparentes al recuperar el foco de la ventana o superar el tiempo de stale, y expone "
      "hooks de mutación (useMutation) que invalidan selectivamente las entradas de caché "
      "afectadas al confirmar un cambio. Por ejemplo, al añadir contenido a una lista se "
      "invalidan simultáneamente las claves lists.mine() y users.lists(), garantizando que "
      "todos los componentes suscritos reflejen el cambio de inmediato.")

    p("Todas las comunicaciones HTTP se canalizan a través de módulos de API dedicados que "
      "emplean la instancia global de Axios preconfigurada con el interceptor de autenticación "
      "y la URL base del servidor.")

    figure("[Figura 4.3: Diagrama de las capas de gestión del estado en el front-end: "
           "Zustand para estado sincrónico y TanStack Query para estado del servidor.]")

    # ── 4.3.2.4 Página de inicio ─────────────────────────────
    h4("Página de inicio")

    p("La página de inicio (HomePage, ruta /) es el punto de entrada principal de la "
      "aplicación. Para usuarios no autenticados muestra un bloque de bienvenida con una "
      "llamada a la acción que invita al registro. Para usuarios autenticados ofrece tres "
      "secciones de contenido personalizado:")

    nid_val = next(nid)
    li("Contenido en tendencia: carrusel horizontal alimentado por el endpoint "
       "GET /api/tmdb/trending, que muestra los títulos cinematográficos y televisivos "
       "más populares del momento.", nid_val)
    li("Recomendaciones personalizadas («Para ti ✨»): carrusel generado a partir del "
       "historial de valoraciones del usuario mediante GET /api/users/me/recommendations. "
       "El algoritmo selecciona los géneros más frecuentes en las valoraciones altas del "
       "usuario y recupera títulos relacionados que aún no han sido valorados.", nid_val)
    li("Usuarios sugeridos: lista de perfiles cinéfilos propuestos para seguir, "
       "fomentando la dimensión social de la plataforma.", nid_val)

    figure("[Figura 4.4: Captura de la página de inicio con las secciones de tendencias, "
           "recomendaciones personalizadas y usuarios sugeridos.]")

    # ── 4.3.2.5 Página de descubrimiento ────────────────────
    h4("Página de descubrimiento")

    p("La página de descubrimiento (DiscoverPage, ruta /discover) centraliza la función de "
      "búsqueda de la aplicación. Dispone de una barra de búsqueda con debounce que lanza "
      "consultas simultáneas a la API de películas y de series, interleaving los resultados "
      "para ofrecer una experiencia de búsqueda mixta sin necesidad de seleccionar un tipo "
      "de contenido. Cuando la barra está vacía se muestran los contenidos en tendencia como "
      "estado por defecto.")

    p("Los resultados se presentan en una cuadrícula de tarjetas (ContentCard) que muestran "
      "el póster, el título y el año de estreno. Cada tarjeta es un enlace a la ficha de "
      "detalle del contenido. La cuadrícula utiliza flex flex-wrap para adaptarse a cualquier "
      "resolución de pantalla sin dejar celdas vacías.")

    figure("[Figura 4.5: Captura de la página de descubrimiento con resultados de búsqueda "
           "mixtos de películas y series.]")

    # ── 4.3.2.6 Ficha de detalle de contenido ───────────────
    h4("Ficha de detalle de contenido")

    p("La ficha de detalle (ContentDetailPage, ruta /content/:contentType/:tmdbId/:slug) "
      "es la vista más compleja de la aplicación y el núcleo de la experiencia de usuario. "
      "Recibe el tipo de contenido (MOVIE o TV) y el identificador de TMDB como parámetros "
      "de ruta. La página se estructura en los siguientes bloques:")

    nid_val = next(nid)
    li("Cabecera (ContentHeader): muestra el backdrop a pantalla completa con degradado "
       "superpuesto, el póster a la izquierda y los metadatos principales a la derecha "
       "(título, año, géneros, puntuación TMDB, número de votos).", nid_val)
    li("Widget de valoración (RatingWidget): permite al usuario puntuar el contenido con "
       "1 a 5 estrellas y adjuntar una reseña textual. Si el usuario ya realizó una "
       "valoración previa, la carga automáticamente para permitir su edición.", nid_val)
    li("Gestión de listas: botón que abre el diálogo AddContentToListDialog, desde el que "
       "el usuario puede añadir el contenido a cualquiera de sus listas (Watchlist, "
       "Favoritos, Visto o listas personalizadas).", nid_val)
    li("Reparto y equipo: galería de tarjetas PersonCard con los principales integrantes, "
       "cada una enlazada a la ficha de la persona (/person/:personId/:slug).", nid_val)
    li("Seguimiento de episodios (solo series TV): acordeón SeasonAccordion que despliega "
       "las temporadas y permite marcar episodios individuales como vistos, registrando "
       "la acción en la entidad EpisodeWatch.", nid_val)

    figure("[Figura 4.6: Captura de la ficha de detalle de una película mostrando el "
           "widget de valoración, los metadatos principales y el carrusel de reparto.]")

    # ── 4.3.2.7 Perfil de usuario ────────────────────────────
    h4("Perfil de usuario")

    p("La página de perfil (ProfilePage, ruta /profile/:username) muestra la información "
      "pública y la actividad de un usuario. Soporta deep-linking mediante el parámetro de "
      "consulta ?tab= para anclar directamente una pestaña concreta. El perfil se organiza "
      "en cuatro pestañas:")

    nid_val = next(nid)
    li("Actividad: línea de tiempo cronológica con las valoraciones, reseñas y movimientos "
       "en listas del usuario, con paginación cliente en bloques de cinco elementos.", nid_val)
    li("Listas: cuadrícula de las listas del usuario con acceso directo a cada "
       "ListDetailPage.", nid_val)
    li("Estadísticas (StatsTab): visualizaciones avanzadas generadas a partir del endpoint "
       "GET /api/users/:id/stats, que devuelve un FullStatsDto con la distribución de "
       "valoraciones (1-5), los géneros más valorados y la actividad mensual.", nid_val)
    li("Insignias (BadgesSection): colección de insignias obtenidas por el usuario a través "
       "del uso de la plataforma, gestionadas por BadgeService con diez tipos definidos "
       "(primera valoración, primera lista, diez valoraciones, etc.).", nid_val)

    p("Para perfiles privados o usuarios no autenticados la página gestiona el código HTTP "
      "403 (ProfilePrivateException) mostrando un mensaje de perfil privado sin reintentar "
      "la petición. Los diálogos FollowListDialog muestran la lista de seguidores y de "
      "cuentas seguidas. El componente QuickEditRatingDialog permite editar una valoración "
      "directamente desde la cuadrícula de pósteres sin abandonar la página.")

    figure("[Figura 4.7: Captura del perfil de usuario mostrando la pestaña de estadísticas "
           "avanzadas con gráficos de distribución de valoraciones y géneros.]")

    # ── 4.3.2.8 Listas personalizadas ───────────────────────
    h4("Listas personalizadas")

    p("La funcionalidad de listas personalizadas se articula en tres páginas "
      "complementarias que comparten la capa de acceso a datos pero difieren en "
      "su presentación y propósito.")

    p("ListsPage (/lists) muestra la cuadrícula de listas del usuario autenticado. Cada "
      "lista se representa mediante el componente MyListCard, que ofrece acceso directo "
      "a las acciones de edición (EditListDialog, para modificar nombre, descripción y "
      "visibilidad) y eliminación. Desde esta misma página se puede crear una nueva lista "
      "personalizada con nombre único por usuario.")

    p("ListDetailPage (/lists/:listId) muestra el contenido completo de una lista concreta "
      "con carga instantánea gracias al patrón placeholderData, que recupera los datos "
      "preliminares del estado de navegación (location.state) mientras se completa la "
      "petición al servidor. En la parte inferior aparece la sección de comentarios, "
      "donde cualquier usuario puede dejar valoraciones públicas sobre la lista "
      "(entidad ListComment, gestionada por ListCommentService).")

    p("SpecialListPage es una vista compartida para las rutas /watchlist, /favorites y "
      "/watched, diferenciadas únicamente por el prop listType (ListType.WATCHLIST, "
      "FAVORITES, WATCHED). El diálogo AddContentToListDialog permite buscar películas y "
      "series directamente desde la interfaz de la lista, deduplicando los resultados "
      "para evitar sugerir contenido ya presente en la misma.")

    figure("[Figura 4.8: Captura de la página de detalle de una lista personalizada "
           "con la cuadrícula de pósteres y la sección de comentarios.]")

    # ── 4.3.2.9 Actividad, notificaciones y tiempo real ─────
    h4("Actividad, notificaciones y comunicación en tiempo real")

    p("La página de actividad (ActivityPage, ruta /activity) muestra un feed cronológico "
      "de las acciones recientes de los usuarios seguidos: nuevas valoraciones, reseñas "
      "publicadas o modificadas, y añadidos a listas. El feed se obtiene del endpoint "
      "GET /api/activity y se presenta con paginación. ActivityService distingue entre "
      "valoraciones nuevas y actualizadas comparando updatedAt con createdAt para generar "
      "los eventos RATING_UPDATED y LIST_UPDATED de forma precisa.")

    p("La página de notificaciones (NotificationsPage, ruta /notifications) lista los "
      "eventos dirigidos al usuario autenticado. El recuento de notificaciones no leídas "
      "se refleja en un badge sobre el icono de la barra de navegación.")

    p("Para la recepción de notificaciones en tiempo real, el componente raíz Layout.tsx "
      "inicializa al cargar la aplicación una conexión WebSocket mediante SockJS y el "
      "protocolo STOMP. El token JWT se incluye en las cabeceras del frame STOMP CONNECT "
      "y es validado por WebSocketAuthInterceptor en el servidor. Cuando se emite un "
      "evento, el cliente actualiza la caché de TanStack Query con setQueryData para "
      "añadir la nueva notificación de forma reactiva e inmediata, sin necesidad de una "
      "nueva petición HTTP, e invalida la consulta para mantener la coherencia a largo plazo.")

    figure("[Figura 4.9: Captura de la bandeja de notificaciones de MovieMate con el "
           "indicador de notificaciones pendientes.]")

    # ── 4.3.2.10 Sistema de diseño y adaptabilidad ──────────
    h4("Sistema de diseño y adaptabilidad")

    p("La interfaz de MovieMate se construye sobre un sistema de diseño coherente definido "
      "mediante tokens de Tailwind CSS 4. Los principales tokens son:")

    nid_val = next(nid)
    li("Paleta oscura estructurada: clases bg-bg-0 a bg-bg-3 para los distintos niveles "
       "de profundidad de los fondos, con gradaciones del azul marino profundo al gris carbón.", nid_val)
    li("Acento dorado (#e8c97a): empleado como color primario de interacción en botones "
       "de acción, estrellas de valoración, iconos activos y elementos de énfasis.", nid_val)
    li("Tipografía dual: Playfair Display para encabezados y títulos de películas, "
       "DM Sans para el cuerpo del texto y controles, DM Mono para fechas y "
       "datos numéricos.", nid_val)

    p("Los componentes accesibles de Radix UI y Shadcn/ui garantizan la compatibilidad "
      "con lectores de pantalla y el cumplimiento de los estándares WCAG. La interfaz es "
      "completamente adaptable (responsive): en pantallas móviles aparece una barra de "
      "navegación inferior fija (BottomNavBar, visible solo en resoluciones inferiores a lg), "
      "mientras que en escritorio se emplea una barra de navegación lateral. Esta "
      "adaptabilidad se logra exclusivamente con Tailwind CSS sin necesidad de librerías "
      "de layout adicionales.")

    figure("[Figura 4.10: Comparativa de la interfaz de MovieMate en escritorio y "
           "dispositivo móvil mostrando la barra de navegación adaptativa.]")

    return paras


def build_branching_strategy_addition(num_ids):
    """Paragraphs to add to §4.5.3 about the branching strategy."""
    nid = iter(num_ids)
    paras = []

    def p(text, style='Normal', num_id=None):
        paras.append(make_para(style, text, num_id=num_id))

    def li(text, nid_val):
        paras.append(make_para('Prrafodelista', text, num_id=nid_val))

    p("Estrategia de ramificación")

    p("El proyecto sigue una estrategia de ramificación (branching strategy) basada en dos "
      "ramas permanentes. La rama develop concentra todo el desarrollo activo: cada nueva "
      "funcionalidad, corrección de errores o tarea de refactorización se implementa en esta "
      "rama antes de ser integrada. Sobre cada push a develop se ejecuta automáticamente el "
      "pipeline de CI/CD, que valida el código mediante pruebas unitarias y análisis estático "
      "antes de construir el artefacto.")

    p("La rama main representa el estado estable y desplegable de la aplicación. Un merge "
      "desde develop a main únicamente se realiza cuando el pipeline completo ha concluido "
      "satisfactoriamente, garantizando que la rama principal nunca contiene código en "
      "estado incorrecto. Un merge exitoso a main desencadena adicionalmente el paso de "
      "construcción y publicación de la imagen Docker, preparando la nueva versión para "
      "su despliegue.")

    p("Esta separación entre rama de integración (develop) y rama de producción (main) "
      "proporciona un cortafuegos de calidad automatizado, elimina la dependencia de "
      "revisiones manuales para detectar regresiones y facilita la trazabilidad entre "
      "versiones desplegadas y commits específicos de la rama principal.")

    return paras


def build_k8s_section(num_ids):
    """Build §4.5.4 Despliegue en Kubernetes con Minikube."""
    nid = iter(num_ids)
    paras = []

    def p(text, style='Normal', num_id=None):
        paras.append(make_para(style, text, num_id=num_id))

    def h3(text):
        paras.append(make_para('Ttulo3', text))

    def li(text, nid_val):
        paras.append(make_para('Prrafodelista', text, num_id=nid_val))

    # Heading
    h3("Despliegue en Kubernetes con Minikube")

    # Intro
    p("Kubernetes es una plataforma de orquestación de contenedores de código abierto que "
      "automatiza el despliegue, el escalado y la gestión de aplicaciones contenerizadas. "
      "Permite describir el estado deseado del sistema mediante ficheros de configuración "
      "declarativos (manifiestos YAML) y se encarga de reconciliar continuamente el estado "
      "real del clúster con el estado deseado.")

    p("Para facilitar la validación local del despliegue orquestado sin necesidad de "
      "infraestructura en la nube, se ha utilizado Minikube. Esta herramienta levanta un "
      "clúster Kubernetes de un único nodo directamente en el equipo del desarrollador, "
      "simulando un entorno de producción real con un coste operativo mínimo.")

    # Manifests
    p("El despliegue de MovieMate en Kubernetes se describe mediante los siguientes recursos:")

    nid_val = next(nid)
    li("Namespace moviemate: agrupa y aísla todos los recursos del proyecto de otros "
       "workloads del clúster, facilitando la gestión de permisos y la limpieza del entorno.", nid_val)
    li("ConfigMap: almacena la configuración no sensible de la aplicación, como el nombre "
       "de la base de datos, el perfil activo de Spring Boot, el puerto de escucha y la "
       "URL del servidor.", nid_val)
    li("Secret: almacena las credenciales sensibles codificadas en base64: contraseña de "
       "PostgreSQL, clave secreta JWT y clave de la API de TMDB. Kubernetes inyecta estos "
       "valores como variables de entorno en los contenedores en tiempo de arranque.", nid_val)
    li("PersistentVolumeClaim (PVC): solicita almacenamiento persistente para los datos "
       "de PostgreSQL, garantizando que la información se conserva entre reinicios de pod "
       "o actualizaciones del Deployment.", nid_val)
    li("Deployment de PostgreSQL: define el pod de base de datos, monta el PVC como "
       "volumen y carga las credenciales desde el Secret.", nid_val)
    li("Service ClusterIP para PostgreSQL: expone la base de datos internamente dentro "
       "del clúster mediante un nombre de servicio estable, desacoplando la dirección IP "
       "del pod de la configuración del back-end.", nid_val)
    li("Deployment del back-end: define el pod de la aplicación Spring Boot con la imagen "
       "Docker generada en el pipeline de CI/CD. Incluye sondas de liveness y readiness "
       "para garantizar que el servicio únicamente recibe tráfico cuando está completamente "
       "iniciado y operativo.", nid_val)
    li("Service NodePort para el back-end: expone el back-end externamente para "
       "permitir el acceso desde el navegador durante el desarrollo local con Minikube.", nid_val)

    # Pipeline steps
    p("El proceso de despliegue local con Minikube sigue los pasos a continuación:")

    nid_val = next(nid)
    li("Inicialización del clúster: minikube start arranca el nodo local y configura "
       "el contexto de kubectl para apuntar al clúster Minikube.", nid_val)
    li("Carga de la imagen: minikube image load moviemate-backend:latest inyecta la "
       "imagen Docker directamente en el registro interno del clúster, evitando la "
       "necesidad de un registro externo en el entorno de desarrollo.", nid_val)
    li("Aplicación de manifiestos: kubectl apply -f k8s/ aplica todos los recursos "
       "definidos en el directorio de manifiestos. Kubernetes gestiona automáticamente "
       "el orden de creación respetando las dependencias entre recursos.", nid_val)
    li("Verificación: kubectl get pods -n moviemate permite monitorizar el estado de "
       "los pods y confirmar que todos se encuentran en estado Running.", nid_val)
    li("Acceso al servicio: minikube service backend-service -n moviemate abre el "
       "servicio en el navegador, resolviendo automáticamente la IP del nodo y el "
       "puerto NodePort asignado.", nid_val)

    # CI/CD integration
    p("En el entorno de integración continua, una vez que el pipeline de GitHub Actions "
      "construye y publica la imagen Docker con éxito tras un merge a main, se ejecuta "
      "automáticamente un paso adicional de despliegue que actualiza el Deployment del "
      "clúster mediante kubectl set image o kubectl apply. Esto garantiza que el clúster "
      "refleja siempre la última versión estable de la aplicación sin intervención manual.")

    p("La combinación de Minikube para el entorno local y Kubernetes para el entorno de "
      "producción garantiza la paridad entre ambos entornos, reduciendo el riesgo de "
      "diferencias de comportamiento entre el desarrollo y la puesta en producción.")

    p("[Figura 4.11: Diagrama del pipeline completo de CI/CD y despliegue automatizado "
      "en Kubernetes con Minikube, desde el push a la rama main hasta el pod en ejecución.]")

    return paras


# ─────────────────────────────────────────────────────────────
#  Main
# ─────────────────────────────────────────────────────────────

def main():
    input_path = "Memoria TFG_rev260318.docx"
    output_path = "Memoria TFG_rev260319.docx"

    doc = Document(input_path)

    # ── Verify paragraph indices before modifying ────────────
    print("Verification:")
    for idx in [414, 415, 421, 486, 508, 509]:
        print(f"  [{idx}] [{doc.paragraphs[idx].style.name}] {doc.paragraphs[idx].text[:80]}")

    # ── Add new numbering definitions (50..62) ────────────────
    for new_id in NEW_NUM_IDS:
        _new_list_num(doc, ABS_NUM_ID, new_id)

    print(f"\nAdded {len(NEW_NUM_IDS)} new list numbering definitions.")

    # ── STEP 1: Add §4.5.4 + branching strategy at end of §4.5 ──
    # Current last para of §4.5.3 is at index 508.
    # (After step 2 the indices will shift, so we work bottom-up.)

    # Re-read indices after no changes yet (safe)
    last_cicd_para = doc.paragraphs[508]

    # Build branching strategy paragraphs (3 paras) → num_ids 60
    branching_paras = build_branching_strategy_addition(NEW_NUM_IDS[10:11])  # numId 60

    # Build K8s section (heading + many paras) → num_ids 61, 62
    k8s_paras = build_k8s_section(NEW_NUM_IDS[11:13])  # numIds 61, 62

    # Insert branching + k8s after last §4.5.3 paragraph
    all_new_end = branching_paras + k8s_paras
    insert_sequence(last_cicd_para._element, all_new_end)
    print(f"Inserted {len(all_new_end)} paragraphs after §4.5.3 end.")

    # ── STEP 2: Replace §4.3.2 content ───────────────────────
    # §4.3.2 heading is at index 414, body paragraphs at 415..421
    # Remove old body (415..421)
    remove_paras(doc, 415, 421)
    print("Removed 7 old §4.3.2 paragraphs.")

    # Insert new content after the heading (now at index 414)
    frontend_paras = build_frontend_section(NEW_NUM_IDS[0:10])  # numIds 50..59
    insert_sequence(doc.paragraphs[414]._element, frontend_paras)
    print(f"Inserted {len(frontend_paras)} paragraphs for §4.3.2.")

    # ── Save ─────────────────────────────────────────────────
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
