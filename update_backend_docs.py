"""
Update Memoria TFG_rev260320.docx with missing backend documentation:
  1. 23 new rows in the endpoints table (Table[3])
  2. 3 new paragraphs in §4.2.1 (7 missing entities)
  3. 3 new paragraphs in §4.3.1 (WebSocket, extra services, Swagger)
→ Saves as Memoria TFG_rev260321.docx
"""

from docx import Document
from docx.oxml import OxmlElement
from lxml import etree
import copy

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def _qn(tag):
    return f'{{{NS}}}{tag}'

def make_para(text, style_id='Normal', bold=False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(_qn('val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        if bold:
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
    return p

def insert_after(ref_elem, items):
    current = ref_elem
    for p_elem in items:
        current.addnext(p_elem)
        current = p_elem
    return current


# ─────────────────────────────────────────────────────────────────────────────
#  Content
# ─────────────────────────────────────────────────────────────────────────────

# New endpoint table rows: (Endpoint, GET, POST, PUT, DELETE, Descripción)
NEW_ENDPOINT_ROWS = [
    # ── UserController ──────────────────────────────────────────────────────
    ('/api/users/me/recommendations', '✔', '',  '',  '',  'Obtener recomendaciones personalizadas'),
    ('/api/users/me/badges',          '✔', '',  '',  '',  'Obtener insignias del usuario actual'),
    ('/api/users/{userId}/badges',    '✔', '',  '',  '',  'Obtener insignias de un usuario'),
    ('/api/users/me/avatar',          '',  '✔', '',  '',  'Subir foto de perfil (multipart/form-data)'),
    # ── ListController ──────────────────────────────────────────────────────
    ('/api/lists/{listId}',           '',  '',  '✔', '✔', 'Editar o eliminar una lista propia'),
    # ── TmdbController ──────────────────────────────────────────────────────
    ('/api/tmdb/trending',            '✔', '',  '',  '',  'Obtener contenido en tendencia (películas + series)'),
    # ── CommentController (comentarios sobre valoraciones) ──────────────────
    ('/api/ratings/{ratingId}/comments',              '✔', '✔', '', '',  'Obtener o crear comentarios de una valoración'),
    ('/api/ratings/{ratingId}/comments/{commentId}',  '',  '',  '', '✔', 'Eliminar un comentario de valoración'),
    # ── ListCommentController ────────────────────────────────────────────────
    ('/api/lists/{listId}/comments',              '✔', '✔', '', '',  'Obtener o crear comentarios de una lista'),
    ('/api/lists/{listId}/comments/{commentId}',  '',  '',  '', '✔', 'Eliminar un comentario de lista'),
    # ── EpisodeWatchController ───────────────────────────────────────────────
    ('/api/episodes/watched/summary',                         '✔', '',  '',  '',  'Progreso de episodios vistos agrupado por serie'),
    ('/api/episodes/watched/{tmdbSeriesId}',                  '✔', '',  '',  '',  'Obtener episodios vistos de una serie'),
    ('/api/episodes/watched/{tmdbId}/{season}/{episode}',     '',  '✔', '',  '',  'Marcar o desmarcar un episodio como visto (toggle)'),
    ('/api/episodes/watched/{tmdbId}/{season}/all',           '',  '✔', '',  '✔', 'Marcar o desmarcar toda una temporada como vista'),
    # ── ReportController ─────────────────────────────────────────────────────
    ('/api/reports',  '', '✔', '', '', 'Crear un reporte de contenido inapropiado'),
    # ── AdminController (requiere rol ADMIN) ─────────────────────────────────
    ('/api/admin/users',            '✔', '',  '',  '',  'Listar todos los usuarios (solo admin)'),
    ('/api/admin/users/{id}/role',  '',  '',  '✔', '',  'Cambiar el rol de un usuario'),
    ('/api/admin/users/{id}/ban',   '',  '',  '✔', '',  'Banear o desbanear un usuario'),
    ('/api/admin/ratings/{id}',     '✔', '',  '',  '✔', 'Ver o eliminar una valoración (moderación)'),
    ('/api/admin/comments/{id}',    '✔', '',  '',  '✔', 'Ver o eliminar un comentario (moderación)'),
    ('/api/admin/reports',                  '✔', '',  '',  '',  'Obtener reportes pendientes de revisión'),
    ('/api/admin/reports/{id}/resolve',     '',  '',  '✔', '',  'Resolver un reporte (elimina el contenido denunciado)'),
    ('/api/admin/reports/{id}/dismiss',     '',  '',  '✔', '',  'Desestimar un reporte sin acción'),
]


# New paragraphs for §4.2.1 Modelo de datos (insert after last paragraph of section)
ENTITY_PARAS = [
    ("Para registrar la interacción de los usuarios sobre las valoraciones, el modelo "
     "incorpora dos entidades adicionales. Comment almacena los comentarios escritos "
     "por otros usuarios sobre una reseña concreta, asociando el texto del comentario "
     "tanto al usuario autor como a la valoración referenciada, con soporte de borrado "
     "lógico (campo deleted). ReviewLike modela los «me gusta» que los usuarios pueden "
     "dar a las reseñas de otros, garantizando mediante una restricción de unicidad que "
     "un mismo usuario no puede dar más de un like a la misma valoración."),

    ("El seguimiento del visionado de contenido televisivo se gestiona a través de la "
     "entidad EpisodeWatch, que almacena los episodios individuales marcados como vistos "
     "por cada usuario. Cada registro identifica la serie mediante su identificador de "
     "TMDB, el número de temporada y el número de episodio, con una restricción de "
     "unicidad sobre la combinación de estos cuatro campos que impide registros "
     "duplicados. Los comentarios sobre las listas personalizadas se modelan mediante "
     "ListComment, que asocia un texto con un usuario autor y la lista comentada, e "
     "incorpora igualmente borrado lógico para preservar la integridad referencial."),

    ("La plataforma incluye además un sistema de moderación de contenido articulado "
     "mediante la entidad ContentReport, que permite a los usuarios reportar valoraciones "
     "o comentarios inapropiados. Cada reporte registra el tipo de objeto denunciado "
     "(RATING o COMMENT), el motivo de la denuncia (SPAM, INAPPROPRIATE, SPOILER u "
     "OTHER) y el estado de gestión (PENDING, RESOLVED o DISMISSED). Finalmente, la "
     "entidad UserBadge materializa el sistema de gamificación de la plataforma, "
     "almacenando las insignias concedidas a cada usuario con la fecha de concesión. "
     "La combinación de usuario e insignia está sujeta a una restricción de unicidad "
     "para garantizar que cada usuario sólo puede obtener cada insignia una vez."),
]


# New paragraphs for §4.3.1 Arquitectura del back-end (insert after figure caption)
ARCH_PARAS = [
    ("El sistema de notificaciones en tiempo real se implementa mediante WebSocket con "
     "el protocolo STOMP sobre SockJS. La clase WebSocketConfig configura el broker de "
     "mensajería y los canales de comunicación, definiendo el prefijo de destino de "
     "aplicación (/app) y el de usuario (/user/queue). La autenticación de las conexiones "
     "WebSocket es gestionada por WebSocketAuthInterceptor, un interceptor personalizado "
     "que extrae el token JWT del frame STOMP CONNECT y lo valida contra el proveedor de "
     "autenticación de Spring Security, asignando el principal autenticado a la sesión "
     "WebSocket antes de que la suscripción al canal privado del usuario sea efectiva."),

    ("Entre los servicios adicionales que complementan la funcionalidad principal del "
     "sistema destacan: BadgeService, responsable de la evaluación y concesión de los "
     "diez tipos de insignias de forma idempotente cada vez que se recalculan las "
     "estadísticas de un usuario; UserStatsService, que mantiene actualizadas las "
     "estadísticas agregadas de cada usuario y genera el FullStatsDto empleado por las "
     "visualizaciones avanzadas del perfil (distribución de valoraciones, géneros más "
     "valorados, actividad mensual); y CacheCleaner, un componente planificado "
     "(@Scheduled) que elimina periódicamente los registros de Content que llevan más "
     "de un umbral configurable de días sin interacción, manteniendo la base de datos "
     "libre de datos obsoletos poco accedidos."),

    ("El back-end incorpora además dos elementos transversales. SwaggerConfig configura "
     "la documentación automática de la API mediante SpringDoc/OpenAPI, generando una "
     "interfaz Swagger UI accesible en /swagger-ui/index.html que describe todos los "
     "endpoints, sus parámetros, cuerpos de petición y respuestas posibles. Por su "
     "parte, @RequirePublicProfile es una anotación personalizada que, combinada con "
     "ProfileSecurityService, permite restringir el acceso a recursos de perfil "
     "lanzando ProfilePrivateException cuando el perfil objetivo es privado y el "
     "solicitante no tiene permisos para consultarlo, simplificando la lógica de "
     "autorización en los controladores correspondientes."),
]


# ─────────────────────────────────────────────────────────────────────────────
#  Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    input_path  = "Memoria TFG_rev260320.docx"
    output_path = "Memoria TFG_rev260321.docx"

    doc = Document(input_path)

    # ── Verify key paragraph indices ─────────────────────────────────────────
    print("Verification:")
    indices_to_check = {
        328: "§4.2.1 last paragraph",
        329: "§4.2.2 heading",
        361: "§4.3.1 heading",
        366: "§4.3.1 figure caption",
        367: "§4.3.1 next heading",
    }
    for idx, label in indices_to_check.items():
        p = doc.paragraphs[idx]
        print(f"  [{idx}] {label}: [{p.style.name}] '{p.text[:70]}'")

    print(f"\nTable 3 before: {len(doc.tables[3].rows)} rows")

    # ── STEP 1: Add rows to endpoints table ──────────────────────────────────
    tbl = doc.tables[3]
    for row_data in NEW_ENDPOINT_ROWS:
        row = tbl.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
    print(f"Table 3 after:  {len(doc.tables[3].rows)} rows (+{len(NEW_ENDPOINT_ROWS)})")

    # ── STEP 2: Add entity paragraphs to §4.3.1 (bottom-up: do this before §4.2.1) ──
    # Insert after paragraph [366] (figure caption at end of §4.3.1)
    ref_4_3_1 = doc.paragraphs[366]._element
    arch_elems = [make_para(t) for t in ARCH_PARAS]
    insert_after(ref_4_3_1, arch_elems)
    print(f"Inserted {len(arch_elems)} paragraphs into §4.3.1.")

    # ── STEP 3: Add entity paragraphs to §4.2.1 ──────────────────────────────
    # Insert after paragraph [328] (last paragraph of §4.2.1)
    ref_4_2_1 = doc.paragraphs[328]._element
    entity_elems = [make_para(t) for t in ENTITY_PARAS]
    insert_after(ref_4_2_1, entity_elems)
    print(f"Inserted {len(entity_elems)} paragraphs into §4.2.1.")

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")

    # Quick spot-check
    print("\nSpot-check §4.2.1 new paragraphs:")
    for i in range(328, 333):
        p = doc.paragraphs[i]
        print(f"  [{i}] {p.text[:100]}")

    print("\nSpot-check §4.3.1 new paragraphs:")
    for i in range(369, 374):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f"  [{i}] {p.text[:100]}")


if __name__ == "__main__":
    main()
