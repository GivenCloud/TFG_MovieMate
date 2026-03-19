"""
Inserts the Extended Abstract (English) into Memoria TFG_rev260319.docx
→ saves as Memoria TFG_rev260320.docx

Structure:
  [94] Título previo  "Extended abstract"   ← keep
  [95] Comentario      editorial note        ← keep
  [96-99] empty Normal paragraphs           ← remove (placeholders)
  → insert 7 sections with bold sub-headings + body paragraphs
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def _qn(tag):
    return f'{{{NS}}}{tag}'

def make_para(text, bold_heading=False):
    """Create a w:p element with Normal style. bold_heading=True makes the run bold."""
    p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(_qn('val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        if bold_heading:
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)

    return p


def insert_sequence(ref_elem, items):
    current = ref_elem
    for p_elem in items:
        current.addnext(p_elem)
        current = p_elem
    return current


def remove_paras(doc, start_idx, end_idx_inclusive):
    body = doc.element.body
    to_remove = [doc.paragraphs[i]._element for i in range(start_idx, end_idx_inclusive + 1)]
    for elem in to_remove:
        body.remove(elem)


# ── Abstract content ─────────────────────────────────────────────────────────

ABSTRACT_PARAS = [

    # ── Context and Motivation ───────────────────────────────────────────────
    ("Context and Motivation", True),

    ("The consumption of audiovisual content has grown exponentially in recent years, "
     "driven by the proliferation of streaming platforms and the widespread availability "
     "of high-speed internet access. In this context, tools dedicated to cataloguing and "
     "rating films and television series — such as Letterboxd, IMDb, or Serializd — have "
     "become indispensable companions for enthusiasts who wish to keep track of their "
     "viewing history, share opinions with other users, and discover new content within a "
     "community. However, these platforms present significant limitations that fragment the "
     "user experience: most specialise in a single content type (either films or series, "
     "but not both simultaneously), lack advanced social features, or do not provide a "
     "unified space in which a user can manage ratings, reviews, and personal lists "
     "together. This situation creates a suboptimal experience for those who consume both "
     "cinema and television and desire a single, integral platform to centralise their "
     "entire audiovisual activity.", False),

    ("MovieMate was conceived to address these shortcomings. It is a full-stack social web "
     "application that combines the content cataloguing capabilities of IMDb, the personal "
     "diary and community features of Letterboxd, and the series tracking functionality of "
     "Serializd into a single, coherent system. The platform allows users to rate and "
     "review both films and television series, organise content into custom lists, follow "
     "other users and track their activity, receive real-time notifications, and obtain "
     "personalised content recommendations based on their viewing history.", False),

    # ── Objectives ───────────────────────────────────────────────────────────
    ("Objectives", True),

    ("The primary objective of this Bachelor's thesis is the design and development of a "
     "responsive, full-stack web application for the consultation, management, and rating "
     "of films and television series. This overarching goal is decomposed into the "
     "following specific sub-objectives: to identify the most appropriate technologies and "
     "tools for the undertaking; to determine the main features and shortcomings of "
     "existing solutions in the domain; to develop an application that addresses those "
     "shortcomings while satisfying the principal quality standards of modern software "
     "engineering; to deploy the application in a production-like environment; and to "
     "document the work carried out thoroughly. It is important to note that the project "
     "aims to deliver a functional first prototype and does not intend to address content "
     "distribution or streaming. External data are sourced from The Movie Database (TMDB) "
     "API to populate the system with real film and series metadata.", False),

    # ── Methodology ──────────────────────────────────────────────────────────
    ("Methodology", True),

    ("The project followed an iterative and incremental development methodology organised "
     "into five sequential phases. The first phase consisted of a preliminary analysis of "
     "existing platforms and an evaluation of candidate technologies for both the "
     "server-side and client-side layers. The second phase covered system analysis and "
     "design: establishing functional and non-functional requirements, designing the "
     "relational data model, specifying the REST API contracts, and producing interface "
     "mockups. The third phase comprised the implementation of both the back-end and the "
     "front-end components. The fourth phase was devoted to verification and validation "
     "through automated unit testing and a comprehensive smoke-test script covering all "
     "API endpoints. The fifth and final phase addressed the documentation of the system "
     "and the preparation of this report.", False),

    ("Throughout development, Git was used for version control following a two-branch "
     "strategy: all active development took place on a dedicated develop branch, and "
     "merges into main were only performed once the automated CI/CD pipeline had completed "
     "all verification steps successfully. This approach ensured that the main branch "
     "always represented a stable, deployable version of the application.", False),

    # ── Back-End ─────────────────────────────────────────────────────────────
    ("Back-End Architecture and Implementation", True),

    ("The server-side component of MovieMate was implemented using Spring Boot 3.3.4 on "
     "Java 21, with PostgreSQL 16 as the relational database management system. The "
     "architecture follows a classic layered pattern divided into four tiers: the "
     "presentation layer (REST controllers), the business logic layer (services), the "
     "persistence layer (Spring Data JPA repositories), and the security layer.", False),

    ("Security is handled by Spring Security in a stateless configuration. Authentication "
     "is based on JSON Web Tokens (JWT): upon successful login, the server issues a signed "
     "token with a 24-hour expiry that the client includes in the Authorization header of "
     "every subsequent request. Passwords are stored using BCrypt hashing. Access control "
     "distinguishes between unauthenticated users, regular authenticated users, and "
     "administrators, with each endpoint protected accordingly.", False),

    ("Content metadata — film and series information, cast, genres, and ratings — is "
     "sourced from The Movie Database API. Rather than pre-loading the entire TMDB "
     "catalogue, MovieMate adopts a lazy synchronisation strategy: content is fetched from "
     "TMDB on demand the first time a user interacts with a title and cached in the local "
     "database. A TTL-based refresh mechanism ensures that frequently accessed or recently "
     "added titles remain up to date, while less popular content is refreshed on a longer "
     "schedule. This approach significantly reduces dependency on the external API and "
     "improves response times for repeated requests.", False),

    ("The data model encompasses fifteen entities: User, Content, Rating, List, "
     "ListContent, ListComment, Follower, FollowRequest, Notification, Activity, Comment, "
     "ReviewLike, UserStats, UserBadge, and EpisodeWatch. The REST API exposes over "
     "seventy endpoints grouped into twelve controllers — authentication, users, content, "
     "ratings, lists, comments, followers, notifications, activity, TMDB, admin, and "
     "episode tracking — all documented via Swagger/OpenAPI.", False),

    ("Real-time notifications are delivered over WebSocket using the STOMP messaging "
     "protocol over SockJS. The server authenticates WebSocket connections by inspecting "
     "the JWT token included in the STOMP CONNECT frame through a custom "
     "WebSocketAuthInterceptor. When an event occurs — a new follower, a like on a review, "
     "or a comment on a list — the server pushes a notification to the corresponding "
     "user's private channel, which the client receives and renders without polling.", False),

    ("Additional back-end features include a gamification subsystem with ten badge types "
     "awarded automatically based on user activity milestones, a personalised "
     "recommendation engine that derives the user's top genre from their rating history "
     "and queries TMDB's discovery endpoint for highly rated titles in that genre, an "
     "episode-level tracking system for television series, a content reporting and "
     "moderation system managed through the admin panel, and a DataSeeder component that "
     "populates a fresh deployment with a realistic set of test users, content, ratings, "
     "and social connections.", False),

    # ── Front-End ─────────────────────────────────────────────────────────────
    ("Front-End Architecture and Implementation", True),

    ("The client-side application was developed as a Single Page Application (SPA) using "
     "React 19 and TypeScript, bundled with Vite 7. The code is organised following a "
     "feature-based directory structure in which each functional module — lists, profile, "
     "discover, notifications, and others — groups its own pages, components, hooks, and "
     "API access logic, minimising coupling between modules.", False),

    ("Global synchronous state — primarily the authenticated user's session data and JWT "
     "token — is managed with Zustand. Server-derived asynchronous state is managed with "
     "TanStack Query v5, which handles the full lifecycle of HTTP requests: caching results "
     "under a structured key hierarchy, automatically revalidating stale data, and "
     "performing selective cache invalidation after mutations. All HTTP communication is "
     "channelled through a global Axios instance configured with a request interceptor "
     "that attaches the JWT Authorization header to every authenticated request.", False),

    ("Navigation is handled by React Router v6 with a centralised route declaration in "
     "App.tsx. Routes are divided into public routes accessible without authentication "
     "and protected routes guarded by a PrivateRoute component that redirects "
     "unauthenticated visitors to the login screen.", False),

    ("The application comprises the following main pages. The home page presents trending "
     "content sourced from TMDB, a personalised recommendations carousel labelled "
     "\"For You\", and a section of suggested users to follow. The discovery page provides "
     "a debounced search bar that queries both films and series simultaneously, "
     "interleaving results by relevance. The content detail page — the most complex view "
     "— displays the backdrop image, poster, metadata, a five-star rating widget that "
     "supports both creation and in-place editing of existing ratings, a dialogue for "
     "adding the title to any personal list, the cast gallery, and for television series "
     "an episode tracking accordion. The profile page is organised into four tabs: "
     "activity timeline, lists, advanced statistics including rating distribution by score "
     "value and top genres, and earned badges. The lists pages allow users to create, "
     "edit, and delete custom lists, browse their contents, and leave public comments. "
     "The activity page displays a chronological feed of actions performed by followed "
     "users. Real-time notifications are reflected immediately in the navigation badge "
     "through the WebSocket connection established by the root layout component.", False),

    ("The visual identity follows a dark theme defined through Tailwind CSS 4 custom "
     "tokens: a deep navy-to-charcoal background scale, a golden accent colour (#e8c97a) "
     "used for all interactive elements, and a dual typography system pairing Playfair "
     "Display for titles with DM Sans for body text. Accessible components from Radix UI "
     "and Shadcn/ui ensure compatibility with assistive technologies. The interface is "
     "fully responsive: a fixed bottom navigation bar replaces the sidebar on mobile "
     "viewports, and all content grids adapt fluidly to any screen size.", False),

    # ── Testing and Deployment ────────────────────────────────────────────────
    ("Testing and Deployment", True),

    ("The quality of the server-side code is validated by a suite of 205 automated unit "
     "tests written with JUnit 5 and Mockito, organised into seventeen test classes that "
     "collectively cover all service layer components. Tests follow the "
     "Arrange-Act-Assert pattern and make extensive use of mock objects to isolate the "
     "unit under test from its dependencies. A complementary shell-based smoke-test script "
     "exercises all API endpoints in an integrated fashion against a running instance of "
     "the application.", False),

    ("The back-end is packaged using a multi-stage Docker build: the first stage compiles "
     "the project with Maven on a Java 21 image and produces the executable JAR; the "
     "second stage copies the JAR into a minimal Eclipse Temurin JRE 21 runtime image, "
     "reducing the final image size and eliminating build-time dependencies. The full "
     "system is orchestrated with Docker Compose, which brings up the back-end and the "
     "PostgreSQL 16 database with a single command. Continuous integration is implemented "
     "via GitHub Actions, triggering on every push to the develop branch and on pull "
     "requests targeting main; it executes the full test suite, builds the Maven artefact, "
     "and verifies code style with Checkstyle. A successful merge to main additionally "
     "triggers the Docker image build job. Kubernetes manifests have been defined for a "
     "production-grade deployment, encompassing a dedicated namespace, ConfigMaps and "
     "Secrets for configuration and credentials, PersistentVolumeClaims for database "
     "storage, Deployments with liveness and readiness probes, and ClusterIP/NodePort "
     "services. Local validation of the Kubernetes configuration is supported through "
     "Minikube, enabling deployment to a single-node cluster on the developer's machine "
     "without cloud infrastructure.", False),

    # ── Results and Conclusions ───────────────────────────────────────────────
    ("Results and Conclusions", True),

    ("MovieMate delivers a fully functional social platform for audiovisual content "
     "management that meets all the objectives established at the outset of the project. "
     "The application successfully integrates content discovery, personal cataloguing, "
     "social interaction, real-time communication, and personalised recommendations in a "
     "single, coherent system accessible from any device. From a technical standpoint, "
     "the project demonstrates the practical application of a broad set of modern web "
     "development technologies and patterns: RESTful API design, stateless JWT-based "
     "security, on-demand external API synchronisation, reactive client-side state "
     "management, WebSocket-based real-time communication, containerised deployment, and "
     "automated CI/CD pipelines. The layered back-end architecture and the feature-based "
     "front-end organisation proved effective choices at this scale, promoting "
     "maintainability and separation of concerns throughout the development lifecycle.", False),

    ("Several avenues for future work have been identified. The recommendation engine "
     "could be enhanced with collaborative filtering algorithms or machine learning "
     "techniques to provide more accurate and personalised suggestions. A native mobile "
     "application for iOS and Android — leveraging React Native to reuse existing client "
     "logic — would significantly broaden the platform's reach. Integration with streaming "
     "service availability APIs would allow users to see, in real time, on which subscribed "
     "platforms a given title can be watched. Finally, at scale, a migration towards a "
     "microservices architecture would enable independent horizontal scaling of "
     "high-demand components such as the notification service and the TMDB synchronisation "
     "pipeline.", False),
]


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    input_path  = "Memoria TFG_rev260319.docx"
    output_path = "Memoria TFG_rev260320.docx"

    doc = Document(input_path)

    print("Verification before edit:")
    for i in range(93, 101):
        print(f"  [{i}] [{doc.paragraphs[i].style.name}] '{doc.paragraphs[i].text[:80]}'")

    # Remove the 4 empty placeholder paragraphs (indices 96-99)
    remove_paras(doc, 96, 99)
    print("\nRemoved 4 empty placeholder paragraphs.")

    # Insert abstract after the comment paragraph (now at index 95)
    ref = doc.paragraphs[95]._element
    elements = [make_para(text, bold) for text, bold in ABSTRACT_PARAS]
    insert_sequence(ref, elements)
    print(f"Inserted {len(elements)} abstract paragraphs.")

    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")

    # Quick verification
    print("\nSpot-check (paragraphs 94-130):")
    for i in range(94, min(130, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f"  [{i}] [{p.style.name}] {p.text[:90]}")


if __name__ == "__main__":
    main()
